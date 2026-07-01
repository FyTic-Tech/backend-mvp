"""Gemini AI client and document processing helpers."""
from __future__ import annotations

import io
import json
import re

from fastapi import HTTPException

from app.config import settings


def _gemini_client():
    if not settings.gemini_api_key:
        raise HTTPException(503, "AI service not configured (GEMINI_API_KEY missing)")
    from google import genai
    return genai.Client(api_key=settings.gemini_api_key)


_GEMINI_MODEL = "gemini-3.5-flash"


def _call(prompt: str) -> str:
    client = _gemini_client()
    response = client.models.generate_content(model=_GEMINI_MODEL, contents=prompt)
    return response.text


def extract_variables(content_blocks: list[str]) -> list[str]:
    """Return unique variable names found in content blocks."""
    found: list[str] = []
    seen: set[str] = set()
    for block in content_blocks:
        for var in re.findall(r"\{\{([A-Z0-9_]+)\}\}", block):
            if var not in seen:
                seen.add(var)
                found.append(var)
    return found


def render_content(raw: list[str], variables: dict) -> list[str]:
    """Substitute {{VAR}} placeholders in content blocks."""
    result = []
    for block in raw:
        for key, value in variables.items():
            block = block.replace(f"{{{{{key}}}}}", value or "")
        result.append(block)
    return result


def compute_progress(
    detected_variables: list[str],
    variables: dict,
    signatories: list[dict],
    signatures: dict,
) -> dict:
    total_vars = len(detected_variables)
    filled_vars = sum(1 for k in detected_variables if variables.get(k, "").strip())
    missing = [k for k in detected_variables if not variables.get(k, "").strip()]
    percent = round(filled_vars / total_vars * 100) if total_vars else 100
    total_sigs = len(signatories)
    filled_sigs = sum(1 for s in signatories if signatures.get(s.get("key", ""), "").strip())
    return {
        "total_vars": total_vars,
        "filled_vars": filled_vars,
        "percent": percent,
        "missing": missing,
        "total_sigs": total_sigs,
        "filled_sigs": filled_sigs,
    }


def extract_text_from_bytes(filename: str, data: bytes) -> str:
    """Extract plain text from any file using markitdown (PDF, DOCX, PPTX, XLSX, etc.)."""
    import os
    import tempfile
    from markitdown import MarkItDown

    ext = ("." + filename.rsplit(".", 1)[-1].lower()) if "." in filename else ".txt"
    if ext == ".txt":
        return data.decode("utf-8", errors="replace")

    with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as f:
        f.write(data)
        tmp = f.name
    try:
        result = MarkItDown().convert(tmp)
        return result.text_content or ""
    finally:
        os.unlink(tmp)


def ai_import_template(raw_text: str, filename: str) -> dict:
    """Send raw document text to Gemini and get back structured template data."""
    prompt = f"""You are a legal document parser for a Mexican law firm app.
Given the following document text, return a JSON object with these fields:
- "name": suggested template name (short, descriptive, in Spanish)
- "content": array of markdown strings (one per paragraph/heading/block)
- "variables": array of unique variable placeholder names you detected or suggest (e.g. "NOMBRE_CLIENTE")
  - Replace variable-looking text like client names, dates, amounts with {{{{VARIABLE_NAME}}}} in the content
  - List those names here
- "signatories": array of objects with "key" (e.g. "CLIENTE"), "label" (display name in Spanish), optional "nameVar"
- "risk_clauses": array of strings describing risky or unusual clauses (in Spanish)

Return ONLY valid JSON. No markdown, no code fences.

Document filename: {filename}
Document text:
{raw_text[:8000]}"""

    raw = _call(prompt)
    raw = raw.strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```[a-z]*\n?", "", raw)
        raw = re.sub(r"```$", "", raw).strip()
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        return {
            "name": filename.rsplit(".", 1)[0],
            "content": [raw_text[:2000]],
            "variables": [],
            "signatories": [],
            "risk_clauses": [],
        }


def ai_summarize(rendered_content: list[str]) -> dict:
    text = "\n".join(rendered_content)
    prompt = f"""Eres un asistente legal para un despacho de abogados en México.
Resume el siguiente contrato en español claro y sencillo.
Devuelve SOLO un JSON con:
- "summary": resumen en 2-4 oraciones
- "key_points": lista de 3-7 puntos clave importantes

Solo JSON, sin markdown.

Contrato:
{text[:6000]}"""
    raw = _call(prompt).strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```[a-z]*\n?", "", raw)
        raw = re.sub(r"```$", "", raw).strip()
    try:
        data = json.loads(raw)
    except json.JSONDecodeError:
        data = {"summary": raw[:500], "key_points": []}
    word_count = len(text.split())
    return {
        "summary": data.get("summary", ""),
        "key_points": data.get("key_points", []),
        "word_count": word_count,
    }


def ai_analyze(rendered_content: list[str]) -> dict:
    text = "\n".join(rendered_content)
    prompt = f"""Eres un abogado senior en México revisando un contrato.
Analiza el contrato y devuelve SOLO un JSON con:
- "sections": array de objetos con "type", "title", "items"
  - Tipos válidos: "parties", "obligations", "risks", "benefits", "clauses"
  - Cada "items" es una lista de strings describiendo los puntos clave

Solo JSON, sin markdown. En español.

Contrato:
{text[:6000]}"""
    raw = _call(prompt).strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```[a-z]*\n?", "", raw)
        raw = re.sub(r"```$", "", raw).strip()
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        return {"sections": [{"type": "clauses", "title": "Análisis", "items": [raw[:300]]}]}


def ai_scan(raw_text: str, filename: str) -> dict:
    prompt = f"""Eres un abogado senior en México revisando un documento legal.
Analiza el siguiente documento y devuelve SOLO un JSON con:
- "markdown": el texto reformateado en markdown limpio (conserva la estructura)
- "analysis": objeto con "sections" (array de {{type, title, items}})
  - Tipos: "parties", "obligations", "risks", "benefits", "clauses"

Solo JSON, sin markdown envolvente.

Documento: {filename}
Texto:
{raw_text[:8000]}"""
    raw = _call(prompt).strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```[a-z]*\n?", "", raw)
        raw = re.sub(r"```$", "", raw).strip()
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        return {
            "markdown": raw_text[:3000],
            "analysis": {"sections": []},
        }


def export_docx(title: str, rendered_content: list[str]) -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(title, level=1)
    for block in rendered_content:
        stripped = block.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        else:
            p = doc.add_paragraph()
            # Handle **bold** and __bold__
            parts = re.split(r"(\*\*.*?\*\*|__.*?__)", stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith("__") and part.endswith("__"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_pdf(title: str, rendered_content: list[str]) -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, topMargin=inch, bottomMargin=inch)
    styles = getSampleStyleSheet()
    story = [Paragraph(title, styles["Title"]), Spacer(1, 0.25 * inch)]

    for block in rendered_content:
        stripped = block.strip()
        if not stripped:
            story.append(Spacer(1, 0.1 * inch))
            continue
        if stripped.startswith("### "):
            story.append(Paragraph(stripped[4:], styles["Heading3"]))
        elif stripped.startswith("## "):
            story.append(Paragraph(stripped[3:], styles["Heading2"]))
        elif stripped.startswith("# "):
            story.append(Paragraph(stripped[2:], styles["Heading1"]))
        else:
            # Convert **bold** / __bold__ to reportlab XML tags
            clean = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", stripped)
            clean = re.sub(r"__(.+?)__", r"<b>\1</b>", clean)
            story.append(Paragraph(clean, styles["Normal"]))

    doc.build(story)
    return buf.getvalue()
